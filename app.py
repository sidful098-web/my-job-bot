import streamlit as st
import pandas as pd
from docx import Document
import google.generativeai as genai
import smtplib
from email.message import EmailMessage
import io
import time

# --- APP CONFIGURATION ---
st.set_page_config(page_title="AI Job Hunter Pro", layout="wide")
st.title("🤖 AI Automated Job Hunter")

# 1. Setup API Keys
with st.sidebar:
    st.header("1. Settings")
    gemini_key = st.text_input("Enter Google Gemini API Key", type="password")
    email_user = st.text_input("Your Gmail Address")
    email_pass = st.text_input("Your Gmail App Password (16 digits)", type="password")
    
if gemini_key:
    try:
        genai.configure(api_key=gemini_key)
        # Verify if model is available
        model = genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        st.sidebar.error(f"API Key Error: {e}")

# 2. File Uploads
st.header("2. Upload Your Files")
col1, col2 = st.columns(2)
with col1:
    master_cv_file = st.file_uploader("Upload Master CV (Word .docx)", type="docx")
with col2:
    job_data = st.file_uploader("Upload HR Excel/Sheet", type=["xlsx", "csv"])

# 3. Main Logic
if st.button("🚀 Start Automated Outreach"):
    if not master_cv_file or not job_data or not gemini_key:
        st.error("❌ Please provide API Key, CV, and Excel file!")
    else:
        try:
            df = pd.read_excel(job_data) if job_data.name.endswith('xlsx') else pd.read_csv(job_data)
            df = df.dropna(subset=['Email']) # Skip rows without emails
            
            for index, row in df.iterrows():
                hr_name = str(row.get('HR Name', 'Hiring Manager'))
                hr_email = str(row.get('Email', '')).strip()
                job_desc = str(row.get('Job Description', 'Position Role'))

                if "@" not in hr_email: continue

                st.info(f"Processing: {hr_name}...")

                try:
                    # AI TAILORING
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    prompt = f"Based on this job: {job_desc}, write a 3-sentence professional summary for a CV."
                    response = model.generate_content(prompt)
                    ai_text = response.text
                    
                    # DOCX CREATION
                    doc = Document(master_cv_file)
                    doc.add_heading('Tailored for this Role', level=1)
                    doc.add_paragraph(ai_text)
                    
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)

                    # EMAIL SENDING
                    msg = EmailMessage()
                    msg['Subject'] = f"Job Application - {hr_name}"
                    msg['From'] = email_user
                    msg['To'] = hr_email
                    msg.set_content(f"Hi {hr_name},\n\nI've attached my tailored CV for your review.\n\nBest regards.")
                    msg.add_attachment(buf.read(), maintype='application', subtype='docx', filename=f"CV_{index}.docx")
                    
                    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
                        smtp.login(email_user, email_pass)
                        smtp.send_message(msg)
                    
                    st.success(f"✅ Sent to {hr_email}")
                    time.sleep(1) 

                except Exception as e:
                    st.error(f"❌ Error at {hr_name}: {str(e)}")
                    
        except Exception as main_e:
            st.error(f"❌ Excel Error: {str(main_e)}")
